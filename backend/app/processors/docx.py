"""
Microsoft Word (.docx) document processor.
"""

from __future__ import annotations

from pathlib import Path
from uuid import UUID

from docx import Document as WordDocument

from app.processors.base import (
    DocumentContent,
    DocumentProcessor,
)


class DocxProcessor(DocumentProcessor):
    """
    Processor for Microsoft Word documents.
    """

    _SUPPORTED_MIME_TYPES = {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }

    _SUPPORTED_EXTENSIONS = {
        ".docx",
    }

    @property
    def supported_mime_types(self) -> set[str]:
        """Return supported MIME types."""
        return self._SUPPORTED_MIME_TYPES

    @property
    def supported_extensions(self) -> set[str]:
        """Return supported file extensions."""
        return self._SUPPORTED_EXTENSIONS

    def process(
        self,
        document_id: UUID,
        path: Path,
    ) -> DocumentContent:
        """
        Extract text from a DOCX document.
        """

        document = WordDocument(path)

        paragraphs = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]

        text = "\n".join(paragraphs)

        core = document.core_properties

        title = (
            core.title.strip()
            if core.title and core.title.strip()
            else path.stem
        )

        metadata = {
            "filename": path.name,
            "extension": path.suffix.lower(),
            "paragraph_count": len(document.paragraphs),
        }

        if core.author:
            metadata["author"] = core.author

        if core.subject:
            metadata["subject"] = core.subject

        if core.category:
            metadata["category"] = core.category

        if core.comments:
            metadata["comments"] = core.comments

        if core.keywords:
            metadata["keywords"] = core.keywords

        if core.language:
            metadata["language"] = core.language

        if core.last_modified_by:
            metadata["last_modified_by"] = core.last_modified_by

        if core.created:
            metadata["created"] = core.created.isoformat()

        if core.modified:
            metadata["modified"] = core.modified.isoformat()

        return DocumentContent(
            document_id=document_id,
            title=title,
            text=text,
            page_count=1,
            metadata=metadata,
        )
